import io

from docx import Document as DocxDocument


class DOCXParser:
    def parse(self, content: bytes) -> str:
        doc = DocxDocument(io.BytesIO(content))
        parts: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style = para.style.name if para.style else ""
            if style.startswith("Heading 1"):
                parts.append(f"# {text}")
            elif style.startswith("Heading 2"):
                parts.append(f"## {text}")
            elif style.startswith("Heading 3"):
                parts.append(f"### {text}")
            elif style.startswith("List"):
                parts.append(f"- {text}")
            else:
                parts.append(text)

        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append("| " + " | ".join(cells) + " |")
            if rows:
                header_sep = "| " + " | ".join("---" for _ in table.rows[0].cells) + " |"
                rows.insert(1, header_sep)
                parts.append("\n".join(rows))

        return "\n\n".join(parts)
